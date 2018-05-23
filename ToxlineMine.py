import urllib2 # For interacting w/ Web API
#import math
#from sys  import argv as argument
from docx import Document # python docx for output
#import re

def expandResults(hits, tempFile): # Helper method meant to increase results size upon user request, only lists 10 results otherwise
    returnIDs = []
    for x in range(0, int(hits)): # for each hit, retrieve reference label from respective web page
        contents = urllib2.urlopen("https://toxgate.nlm.nih.gov/cgi-bin/sis/search2/g?" + tempFile + ":" +str(x)).read()
        IDStr = contents[contents.find("<Id>")+4:contents.find("</Id>")+(-1)]
        IDsTemp = IDStr.split(" ")
        for label in IDsTemp:
            returnIDs.append(label)
    #remove duplicates from list
    finalIDs = []
    for label in returnIDs:
        if label not in finalIDs:
            finalIDs.append(label)
    #return statement
    #returnIDs.sort()
    return finalIDs

print('This script will query and return based on keywords entered....')
keyword = raw_input("Enter search words: ")

# Making list of keywords in order to bold in results later
keywordTemp = keyword
keywordTemp = keywordTemp.lower() # For case-insensitive comparison later....
keywordTemp = keywordTemp.replace("(","") # Removing special characters/words
keywordTemp = keywordTemp.replace(")","")
keywordTemp = keywordTemp.replace("and","")
keywordTemp = keywordTemp.replace("or","")
keywordTemp = keywordTemp.replace("not","")
keyList = keywordTemp.split(" ")
while '' in keyList:
    keyList.remove('')

document = Document() # Using python-docx to create output file

p = document.add_paragraph('')           # Creating header for document of results
p.add_run('Toxline Search Query: ').bold = True
p.add_run(keyword)         # Outputing the query to the output file

keyword = keyword.replace(" ","+")
contents = urllib2.urlopen("https://toxgate.nlm.nih.gov/cgi-bin/sis/search2/x?dbs+toxline:" + keyword).read()

hits = contents[contents.find("<Count>")+7:contents.find("</Count>")]
p = document.add_paragraph('')
p.add_run('Hits: ').bold = True
p.add_run(hits)

tempFile = contents[contents.find("<TemporaryFile>")+15:contents.find("</TemporaryFile>")]
#print(tempFile + " TempLink")

answer = raw_input(hits + ' results were found. Do you want to list them all? (y/n)')
if(answer == "y"):
    IDs = expandResults(hits, tempFile) # If "y" then lists all results otherwise it will list 10
else:
    IDStr = contents[contents.find("<Id>")+4:contents.find("</Id>")+(-1)]
    IDs = IDStr.split(" ")
#print(IDs)

drugName = raw_input("What drug did you want?: ") # Specifies drugname to filter text, database search query might not include drugname
count = 0

p = document.add_paragraph()
p.add_run('Drug: ').bold = True
p.add_run(drugName.capitalize())

# Printing the desired contents of the file
for label in IDs: # Uses Web API to pull data from Toxline DB based on each article
    data = urllib2.urlopen("https://toxgate.nlm.nih.gov/cgi-bin/sis/search2/z?dbs+toxline:@term+@DOCNO+" + label).read()
    if(drugName in data or drugName.capitalize() in data):
        count = count + 1
        if "<na>" in data: # prints title of source
            title = data[data.find("<na>")+4:data.find("</na>")]
            title = title.strip() # Will remove \n if present
            p = document.add_paragraph()
            p.add_run(title).bold = True
        if "<au>" in data: # prints author
            authors = data[data.find("<au>")+4:data.find("</au>")]
            authors = authors.strip() # Will remove \n if present
            p = document.add_paragraph()
            p.add_run(authors)
        if "<so>" in data: # prints source
            source =  data[data.find("<so>")+4:data.find("</so>")]
            source = source.strip() # Will remove \n if present
            p = document.add_paragraph()
            p.add_run(source)
        if "<ab>" in data: # prints abstract (the important thing)
            abstract = data[data.find("<ab>")+4:data.find("</ab>")]
            abstract = abstract.strip() # Will remove \n if present
            p = document.add_paragraph()
            for word in abstract.split(): # Searches for keywords in abstract in order to bold for output
                foundWord = None
                for key in keyList:
                    if(key in word.lower()):
                        foundWord = key
                if foundWord is not None:
                    #print(foundWord)
                    keyPos = word.lower().find(foundWord)
                    p.add_run(word[0:keyPos])
                    p.add_run(word[keyPos:keyPos+len(foundWord)]).bold = True
                    p.add_run(word[keyPos+len(foundWord):] + ' ')
                else:
                    p.add_run(word + ' ')
        p = document.add_paragraph(' ')

document.save(drugName.capitalize() + 'ToxlineSearch.docx') # Saves document and prints end to user
print(str(count) + " articles were found with the drug. See created document..." )
